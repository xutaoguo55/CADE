#!/usr/bin/env python3
"""Build NAR Ge&B manuscript DOCX with embedded figures.

Produces a complete DOCX where all 7 main figures and 7 supplementary figures
are embedded at the end of the file, after the figure legends, as is the
standard for NAR Ge&B submissions.
"""

from __future__ import annotations

import shutil
import subprocess
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
SOURCE_MD = MANUSCRIPT_DIR / "CADE_NARGeB_manuscript.md"
DOCX_OUT = MANUSCRIPT_DIR / "CADE_NARGeB_manuscript.docx"
FIGURES_DIR = ROOT / "figures"
EMBED_PNG_DIR = MANUSCRIPT_DIR / "docx_embedded_figures_png"

FIGURE_INSERTS = {
    "Figure 1.": "Figure1_Workflow_Summary.png",
    "Figure 2.": ["Figure2_CADE_Benchmark.png", "Figure2E_5method_AUROC.png"],
    "Figure 3.": "Figure3_CADE_FHL_Multi_Panel.png",
    "Figure 4.": "Figure4_Composition_Correlation.png",
    "Figure 5.": "Figure5_Pathway_Analysis.png",
    "Figure 6.": "Figure6_WGCNA_Modules.png",
    "Figure 7.": "Figure7_TF_TargetGene_Signature.png",
    "SuppFig S1.": "SuppFigure_S1_CrossDisease_Heatmap.png",
    "SuppFig S2.": "SuppFigure_S2_Sepsis_Validation.png",
    "SuppFig S3.": "SuppFigure_S3_MAS_Validation.png",
    "SuppFig S4.": "SuppFigure_S4_Sensitivity_Analyses.png",
    "SuppFig S5.": "SuppFigure_S5_Extended_Benchmark.png",
    "SuppFig S6.": "SuppFigure_S6_Parameter_Sensitivity.png",
    "SuppFig S7.": "SuppFigure_S7_CADE_ILR_RankStability.png",
}


def ensure_pngs():
    """Convert each TIF to PNG for DOCX embedding."""
    EMBED_PNG_DIR.mkdir(exist_ok=True)
    from PIL import Image
    for png_name in set(
        png
        for v in FIGURE_INSERTS.values()
        for png in (v if isinstance(v, list) else [v])
    ):
        png_path = EMBED_PNG_DIR / png_name
        if png_path.exists():
            continue
        tif_path = FIGURES_DIR / png_name.replace(".png", ".tif")
        if not tif_path.exists():
            print(f"WARNING: {tif_path} not found")
            continue
        Image.open(tif_path).save(png_path, "PNG", optimize=True)
        print(f"  Generated: {png_path.name}")


def build():
    ensure_pngs()
    # Generate base DOCX
    subprocess.run(
        ["pandoc", str(SOURCE_MD), "-o", str(DOCX_OUT)],
        check=True,
    )
    # Append figure references to the document
    append_figures()
    print(f"Wrote {DOCX_OUT}")


def append_figures():
    """Append a 'Figure appendix' section with all embedded images to the DOCX."""
    # Use python-docx to add images at end
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed; installing...")
        subprocess.run(["pip3", "install", "python-docx", "--quiet"], check=True)
        from docx import Document

    doc = Document(str(DOCX_OUT))

    # Page break before figure appendix
    doc.add_page_break()
    h = doc.add_heading("Figure Appendix (embedded images)", level=1)

    main_figs = ["Figure 1.", "Figure 2.", "Figure 3.", "Figure 4.", "Figure 5.", "Figure 6.", "Figure 7."]
    supp_figs = ["SuppFig S1.", "SuppFig S2.", "SuppFig S3.", "SuppFig S4.", "SuppFig S5.", "SuppFig S6.", "SuppFig S7."]

    for prefix in main_figs + supp_figs:
        pngs = FIGURE_INSERTS.get(prefix, [])
        if isinstance(pngs, str):
            pngs = [pngs]
        for png in pngs:
            png_path = EMBED_PNG_DIR / png
            if not png_path.exists():
                print(f"  Missing PNG: {png_path}")
                continue
            doc.add_paragraph(f"\n[{prefix} {png.replace('_', ' ').replace('.png', '')}]")
            try:
                doc.add_picture(str(png_path), width=None)
            except Exception as e:
                print(f"  Error adding {png}: {e}")

    doc.save(str(DOCX_OUT))
    print(f"  Embedded {len([p for v in FIGURE_INSERTS.values() for p in (v if isinstance(v, list) else [v])])} figures")


if __name__ == "__main__":
    build()
